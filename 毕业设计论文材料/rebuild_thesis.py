"""
重建论文Word文档：
- 嵌入所有必要图表（架构图、ER图、类图、时序图、活动图、代码截图、项目截图）
- 格式：小四号字体(12pt)、1.25倍行距、无段前段后间距
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE  = r"E:\AI\house\project\daojishi\毕业设计论文材料"
FIGS  = os.path.join(BASE, "figures_generated")
PPT   = os.path.join(BASE, "..", "docs", "ppt_export_after")

IMG = {
    "fig4-1":      os.path.join(FIGS, "fig4-1-architecture.png"),
    "fig5-1":      os.path.join(FIGS, "fig5-1-er.png"),
    "fig5-2":      os.path.join(FIGS, "fig5-2-class.png"),
    "fig5-3":      os.path.join(FIGS, "fig5-3-sequence.png"),
    "fig5-4":      os.path.join(FIGS, "fig5-4-activity.png"),
    "code-auth":   os.path.join(FIGS, "code-auth.png"),
    "code-payment":os.path.join(FIGS, "code-payment.png"),
    "screenshot1": os.path.join(PPT,  "幻灯片3.PNG"),
    "screenshot2": os.path.join(PPT,  "幻灯片4.PNG"),
    "screenshot3": os.path.join(PPT,  "幻灯片5.PNG"),
    "screenshot4": os.path.join(PPT,  "幻灯片6.PNG"),
}

MD_FILE  = os.path.join(BASE, "毕业设计论文_完善版.md")
OUT_FILE = os.path.join(BASE, "基于 Spring Boot 的房屋交易平台设计与实现-完善版.docx")

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_run_font(run, size_pt=12, bold=False, font_name="宋体"):
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.name = font_name
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)

def set_para_format(para, spacing_line=276, space_before=0, space_after=0):
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:line"), str(spacing_line))
    spacing.set(qn("w:lineRule"), "auto")
    spacing.set(qn("w:before"), str(space_before))
    spacing.set(qn("w:after"), str(space_after))

def add_heading(doc, text, level):
    para = doc.add_paragraph()
    set_para_format(para, spacing_line=312, space_before=240, space_after=120)
    run = para.add_run(text)
    if level == 1:
        set_run_font(run, size_pt=16, bold=True, font_name="黑体")
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        set_run_font(run, size_pt=14, bold=True, font_name="黑体")
    else:
        set_run_font(run, size_pt=12, bold=True, font_name="黑体")
    return para

def add_body_para(doc, text):
    para = doc.add_paragraph()
    set_para_format(para)
    pPr = para._p.get_or_add_pPr()
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind"); pPr.append(ind)
    ind.set(qn("w:firstLineChars"), "200")
    ind.set(qn("w:firstLine"), "240")
    run = para.add_run(text)
    set_run_font(run, size_pt=12)
    return para

def add_image_para(doc, img_key, caption):
    para = doc.add_paragraph()
    set_para_format(para, space_before=60, space_after=60)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_path = IMG.get(img_key)
    if img_path and os.path.exists(img_path):
        run = para.add_run()
        run.add_picture(img_path, width=Cm(14))
        print(f"  ✔ {img_key}")
    else:
        run = para.add_run(f"[图片未找到: {img_key}]")
        set_run_font(run, size_pt=10)
        print(f"  ✘ NOT FOUND: {img_key} -> {img_path}")
    cap = doc.add_paragraph()
    set_para_format(cap, space_before=0, space_after=120)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = cap.add_run(caption)
    set_run_font(run2, size_pt=10.5, font_name="宋体")

IMG_PATTERN = re.compile(r'!\[([^\]]+)\]\(([^)]+)\)')

def path_to_key(path):
    p = path.replace("\\", "/")
    if "fig4-1" in p: return "fig4-1"
    if "fig5-1" in p: return "fig5-1"
    if "fig5-2" in p: return "fig5-2"
    if "fig5-3" in p: return "fig5-3"
    if "fig5-4" in p: return "fig5-4"
    if "code-auth" in p: return "code-auth"
    if "code-payment" in p: return "code-payment"
    if "幻灯片5" in p: return "screenshot1"
    if "幻灯片7" in p: return "screenshot2"
    return None

# ─── Build ────────────────────────────────────────────────────────────────────

def build():
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.17)
        section.right_margin  = Cm(3.17)

    with open(MD_FILE, encoding="utf-8") as f:
        lines = f.readlines()

    print("Building thesis document...")

    # ── 表格缓冲 ──────────────────────────────────────────────────
    table_rows = []
    in_table = False

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False; return
        # remove separator rows (only dashes)
        data = [r for r in table_rows if not re.match(r'^\|[-| :]+\|$', r.strip())]
        if not data:
            table_rows = []; in_table = False; return
        cols = [c.strip() for c in data[0].strip().strip('|').split('|')]
        ncol = len(cols)
        tbl = doc.add_table(rows=len(data), cols=ncol)
        tbl.style = 'Table Grid'
        for ri, row_str in enumerate(data):
            cells = [c.strip() for c in row_str.strip().strip('|').split('|')]
            for ci in range(ncol):
                txt = cells[ci] if ci < len(cells) else ''
                cell = tbl.cell(ri, ci)
                para = cell.paragraphs[0]
                run = para.add_run(txt)
                set_run_font(run, size_pt=10.5, bold=(ri == 0))
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_rows = []; in_table = False
    # ──────────────────────────────────────────────────────────────

    for raw in lines:
        line = raw.rstrip("\n")

        # 表格行检测
        if line.strip().startswith('|') and line.strip().endswith('|'):
            in_table = True
            table_rows.append(line)
            continue
        else:
            if in_table:
                flush_table()

        if not line.strip():
            continue

        m = IMG_PATTERN.match(line.strip())
        if m:
            caption = m.group(1); path = m.group(2)
            key = path_to_key(path)
            add_image_para(doc, key, caption)
            if key == "screenshot1":
                add_image_para(doc, "screenshot2", "图5-8 系统功能展示二（房源列表与筛选）")
                add_image_para(doc, "screenshot3", "图5-9 系统功能展示三（管理员后台）")
                add_image_para(doc, "screenshot4", "图5-10 系统功能展示四（经纪人工作台）")
            continue

        # 标题行
        if line.startswith("#### "):
            add_heading(doc, line[5:], 3)
        elif line.startswith("### "):
            add_heading(doc, line[4:], 3)
        elif line.startswith("## "):
            add_heading(doc, line[3:], 2)
        elif line.startswith("# "):
            add_heading(doc, line[2:], 1)
        # 参考文献条目
        elif re.match(r'^\[\d+\]', line.strip()):
            para = doc.add_paragraph()
            set_para_format(para)
            run = para.add_run(line.strip())
            set_run_font(run, size_pt=12)
        # 表格说明行（以"表"开头且不是正文）
        elif re.match(r'^表\s*\d+', line.strip()):
            para = doc.add_paragraph()
            set_para_format(para, space_before=60, space_after=60)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line.strip())
            set_run_font(run, size_pt=10.5, bold=True)
        # **加粗文本** 段落
        elif line.strip().startswith('**') and '：' in line:
            para = doc.add_paragraph()
            set_para_format(para)
            pPr = para._p.get_or_add_pPr()
            ind = pPr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind"); pPr.append(ind)
            ind.set(qn("w:firstLineChars"), "200")
            ind.set(qn("w:firstLine"), "240")
            # parse bold prefix
            txt = line.strip()
            bold_match = re.match(r'\*\*(.+?)\*\*(.+)', txt)
            if bold_match:
                r1 = para.add_run(bold_match.group(1))
                set_run_font(r1, size_pt=12, bold=True)
                r2 = para.add_run(bold_match.group(2))
                set_run_font(r2, size_pt=12)
            else:
                run = para.add_run(txt.replace('**',''))
                set_run_font(run, size_pt=12)
        # 普通正文
        else:
            txt = line.strip()
            # strip inline markdown bold
            txt = re.sub(r'\*\*(.+?)\*\*', r'\1', txt)
            # strip inline code backticks
            txt = re.sub(r'`([^`]+)`', r'\1', txt)
            if txt:
                add_body_para(doc, txt)

    # 最后如果还有未刷新的表格
    if in_table:
        flush_table()

    doc.save(OUT_FILE)
    print(f"\n✔ Saved: {OUT_FILE}")

if __name__ == "__main__":
    build()

