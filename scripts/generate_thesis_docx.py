from __future__ import annotations

from pathlib import Path
import re
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
THESIS_MD = ROOT / "毕业设计论文材料" / "毕业设计论文_最终版.md"
OUT_DIR = ROOT / "毕业设计论文材料"
GENERATED_DIR = OUT_DIR / "generated"
OUT_DOCX = OUT_DIR / "基于 Spring Boot 的房屋交易平台设计与实现-最终版.docx"


def ensure_dir(path: Path) -> None:
    path.mkdir(parents=True, exist_ok=True)


def load_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="ignore")


def find_font(candidates: Iterable[str]) -> str:
    for candidate in candidates:
        if Path(candidate).exists():
            return candidate
    return "C:/Windows/Fonts/simsun.ttc"


def render_code_image(src: Path, dst: Path, title: str, start_marker: str, end_marker: str | None = None) -> None:
    code = load_text(src)
    start = code.find(start_marker)
    if start == -1:
        snippet = code[:1600]
    else:
        end = code.find(end_marker, start + len(start_marker)) if end_marker else -1
        snippet = code[start:end] if end != -1 else code[start:start + 1800]
    snippet = snippet.strip().replace("\t", "    ")

    font_path = find_font([
        "C:/Windows/Fonts/consola.ttf",
        "C:/Windows/Fonts/consolab.ttf",
        "C:/Windows/Fonts/simfang.ttf",
    ])
    title_font_path = find_font([
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
    ])
    code_font = ImageFont.truetype(font_path, 26)
    title_font = ImageFont.truetype(title_font_path, 30)

    lines = snippet.splitlines()
    pad = 36
    line_h = 40
    width = 1700
    height = pad * 3 + 60 + line_h * max(len(lines), 8)
    image = Image.new("RGB", (width, height), "#f7f8fb")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((20, 20, width - 20, height - 20), radius=24, outline="#c8d3e1", width=3, fill="#ffffff")
    draw.rounded_rectangle((40, 40, width - 40, 120), radius=18, outline="#dae3ee", width=2, fill="#eef4fb")
    draw.text((64, 62), title, font=title_font, fill="#183153")
    draw.rounded_rectangle((40, 140, width - 40, height - 40), radius=18, outline="#dde5ef", width=2, fill="#fbfcfe")

    y = 170
    for idx, line in enumerate(lines, start=1):
        num = f"{idx:>2}  "
        draw.text((70, y), num, font=code_font, fill="#8a97a8")
        draw.text((140, y), line[:100], font=code_font, fill="#233142")
        y += line_h
        if y > height - 80:
            break

    ensure_dir(dst.parent)
    image.save(dst)


def set_paragraph_format(paragraph, first_line_indent: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent:
        fmt.first_line_indent = Cm(0.74)


def set_run_font(run, east_asia: str = "宋体", ascii_font: str = "Times New Roman", size: int = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.name = ascii_font
    run.font.size = Pt(size)
    rpr = run._element.rPr
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        run._element.insert(0, rpr)
    fonts = rpr.find(qn("w:rFonts"))
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        rpr.append(fonts)
    fonts.set(qn("w:ascii"), ascii_font)
    fonts.set(qn("w:hAnsi"), ascii_font)
    fonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(document: Document, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=False, size=12, east_asia="宋体", indent=True):
    p = document.add_paragraph()
    p.alignment = align
    set_paragraph_format(p, first_line_indent=indent)
    run = p.add_run(text)
    set_run_font(run, east_asia=east_asia, size=size, bold=bold)
    return p


def add_center_caption(document: Document, text: str):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line_indent=False)
    run = p.add_run(text)
    set_run_font(run, east_asia="宋体", size=11)


def parse_markdown_to_doc(document: Document, markdown: str):
    image_pattern = re.compile(r"^!\[(.+?)\]\((.+?)\)$")

    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            continue

        if line.startswith("# "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p, first_line_indent=False)
            run = p.add_run(line[2:].strip())
            set_run_font(run, east_asia="黑体", size=16, bold=True)
            continue

        if line.startswith("## "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, first_line_indent=False)
            run = p.add_run(line[3:].strip())
            set_run_font(run, east_asia="黑体", size=14, bold=True)
            continue

        if line.startswith("### "):
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph_format(p, first_line_indent=False)
            run = p.add_run(line[4:].strip())
            set_run_font(run, east_asia="黑体", size=12, bold=True)
            continue

        img_match = image_pattern.match(line.strip())
        if img_match:
            caption, img_path = img_match.groups()
            img_path = Path(img_path)
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_format(p, first_line_indent=False)
            run = p.add_run()
            run.add_picture(str(img_path), width=Cm(15.8))
            add_center_caption(document, caption)
            continue

        if re.match(r"^\[\d+\]", line):
            add_paragraph(document, line, indent=False)
            continue

        if line.startswith("- "):
            add_paragraph(document, "• " + line[2:].strip(), indent=False)
            continue

        if re.match(r"^\d+\.\s", line):
            add_paragraph(document, line.strip(), indent=False)
            continue

        add_paragraph(document, line.strip())


def add_cover(document: Document):
    section = document.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("本科毕业论文")
    set_run_font(run, east_asia="黑体", size=22, bold=True)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Spring Boot 的房屋交易平台设计与实现")
    set_run_font(run, east_asia="黑体", size=18, bold=True)

    for _ in range(3):
        document.add_paragraph()

    table = document.add_table(rows=6, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("学 校", "湖南科技大学"),
        ("学 院", "计算机科学与工程学院"),
        ("专 业", "数据科学与大数据技术"),
        ("学生姓名", "刘思博"),
        ("学 号", "2205060209"),
        ("论文时间", "2026 年 4 月"),
    ]
    for idx, (k, v) in enumerate(rows):
        for text, cell in zip((k, v), table.rows[idx].cells):
            cell.text = text
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_format(p, first_line_indent=False)
                for r in p.runs:
                    set_run_font(r, east_asia="宋体", size=12, bold=(cell is table.rows[idx].cells[0]))

    document.add_page_break()


def extract_outline(markdown: str) -> list[tuple[int, str]]:
    outline: list[tuple[int, str]] = []
    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if line.startswith("## ") or line.startswith("### "):
            level = 2 if line.startswith("## ") else 3
            outline.append((level, line.split(" ", 1)[1].strip()))
    return outline


def add_manual_toc(document: Document, outline: list[tuple[int, str]]):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p, first_line_indent=False)
    run = p.add_run("目  录")
    set_run_font(run, east_asia="黑体", size=16, bold=True)

    for level, text in outline:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_paragraph_format(p, first_line_indent=False)
        if level == 3:
            p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(text)
        set_run_font(run, east_asia="宋体", size=12)

    document.add_page_break()


def main():
    ensure_dir(GENERATED_DIR)
    render_code_image(
        ROOT / "backend" / "src" / "main" / "java" / "com" / "trae" / "housing" / "controller" / "AuthController.java",
        GENERATED_DIR / "code-auth.png",
        "AuthController 关键代码截图",
        "@PostMapping(\"/register\")",
        "@PostMapping(\"/forgot-password\")",
    )
    render_code_image(
        ROOT / "backend" / "src" / "main" / "java" / "com" / "trae" / "housing" / "controller" / "PaymentController.java",
        GENERATED_DIR / "code-payment.png",
        "PaymentController 关键代码截图",
        "@PostMapping(\"/orders\")",
        "private User getCurrentUser()",
    )

    markdown = load_text(THESIS_MD)
    markdown_lines = markdown.splitlines()
    if markdown_lines and markdown_lines[0].startswith("# "):
        markdown = "\n".join(markdown_lines[1:]).lstrip()
    outline = extract_outline(markdown)

    document = Document()
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    add_cover(document)
    add_manual_toc(document, outline)
    parse_markdown_to_doc(document, markdown)
    document.save(OUT_DOCX)
    print(f"generated: {OUT_DOCX}")


if __name__ == "__main__":
    main()
